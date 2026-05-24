from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "memoria_proyecto_intermodular.md"
TARGET = ROOT / "docs" / "memoria_proyecto_intermodular.docx"


def set_document_properties(doc):
    props = doc.core_properties
    props.title = "Sistema de recomendación de riego con AEMET y Machine Learning"
    props.subject = "Proyecto intermodular de agricultura de precisión"
    props.author = "Román Maeztu"
    props.last_modified_by = "Román Maeztu"
    props.comments = ""
    props.keywords = "riego, AEMET, Machine Learning, agricultura de precisión"
    props.category = "Proyecto intermodular"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def clean_inline(text):
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"`([^`]*)`", r"\1", text)
    return text.replace("  ", " ").strip()


def add_markdown_paragraph(doc, text, style=None):
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.space_after = Pt(6)

    parts = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        else:
            paragraph.add_run(part)
    return paragraph


def apply_document_styles(doc):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(22)
    title.font.bold = True
    title.font.color.rgb = RGBColor(11, 37, 69)
    title.paragraph_format.space_after = Pt(10)

    for name, size, color in [
        ("Heading 1", 16, RGBColor(31, 78, 121)),
        ("Heading 2", 14, RGBColor(31, 78, 121)),
        ("Heading 3", 12, RGBColor(31, 78, 121)),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)


def add_table(doc, rows):
    if not rows:
        return
    table = doc.add_table(rows=0, cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True

    for row_index, row_values in enumerate(rows):
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            cells[idx].text = clean_inline(value)
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cells[idx])
            for paragraph in cells[idx].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(9)
            if row_index == 0:
                set_cell_shading(cells[idx], "E8EEF5")
                for paragraph in cells[idx].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

    doc.add_paragraph()


def convert_markdown_to_docx():
    doc = Document()
    set_document_properties(doc)
    apply_document_styles(doc)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    pending_table = []
    in_code_block = False
    code_lines = []
    first_heading = True

    def flush_table():
        nonlocal pending_table
        if pending_table:
            add_table(doc, pending_table)
            pending_table = []

    def flush_code():
        nonlocal code_lines
        if code_lines:
            paragraph = doc.add_paragraph()
            paragraph.style = doc.styles["Normal"]
            run = paragraph.add_run("\n".join(code_lines))
            run.font.name = "Consolas"
            run.font.size = Pt(8)
            code_lines = []

    for raw_line in lines:
        line = raw_line.rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code_block:
                flush_code()
                in_code_block = False
            else:
                flush_table()
                in_code_block = True
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        if not stripped:
            flush_table()
            continue

        image_match = re.match(r"!\[(.*?)\]\((.*?)\)", stripped)
        if image_match:
            flush_table()
            image_path = (SOURCE.parent / image_match.group(2)).resolve()
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if image_path.exists():
                paragraph.add_run().add_picture(str(image_path), width=Inches(6.2))
            else:
                paragraph.add_run(f"[Imagen no encontrada: {image_path}]")
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            cells = [cell.strip() for cell in stripped.strip("|").split("|")]
            if all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
                continue
            pending_table.append(cells)
            continue

        flush_table()

        if stripped.startswith("# "):
            text = clean_inline(stripped[2:])
            paragraph = doc.add_paragraph(text, style="Title" if first_heading else "Heading 1")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if first_heading else WD_ALIGN_PARAGRAPH.LEFT
            first_heading = False
        elif stripped.startswith("## "):
            doc.add_paragraph(clean_inline(stripped[3:]), style="Heading 1")
        elif stripped.startswith("### "):
            doc.add_paragraph(clean_inline(stripped[4:]), style="Heading 2")
        elif stripped.startswith("#### "):
            doc.add_paragraph(clean_inline(stripped[5:]), style="Heading 3")
        elif stripped.startswith("> "):
            paragraph = add_markdown_paragraph(doc, clean_inline(stripped[2:]))
            paragraph.paragraph_format.left_indent = Cm(0.5)
        elif re.match(r"^\d+\.\s+", stripped):
            paragraph = add_markdown_paragraph(doc, re.sub(r"^\d+\.\s+", "", stripped), style="List Number")
            paragraph.paragraph_format.space_after = Pt(3)
        elif stripped.startswith("- "):
            paragraph = add_markdown_paragraph(doc, stripped[2:], style="List Bullet")
            paragraph.paragraph_format.space_after = Pt(3)
        elif re.match(r"^\*\*Figura \d+\.", stripped):
            paragraph = add_markdown_paragraph(doc, stripped)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.keep_with_next = True
        else:
            add_markdown_paragraph(doc, stripped)

    flush_table()
    flush_code()

    doc.save(TARGET)
    print(TARGET)


if __name__ == "__main__":
    convert_markdown_to_docx()
